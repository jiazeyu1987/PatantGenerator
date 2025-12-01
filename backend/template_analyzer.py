"""
专利模板智能分析器

负责深度分析 DOCX 模板文件，提取结构、格式、内容要求等信息，
为专利生成智能体提供详细的模板指导信息。
"""

import re
import logging
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logger = logging.getLogger(__name__)


@dataclass
class TemplateStructure:
    """模板结构信息"""
    hierarchy: Dict[str, Any]  # 标题层级结构
    sections: List[Dict[str, Any]]  # 章节详细信息
    section_count: int  # 章节总数
    max_heading_level: int  # 最大标题级别
    placeholder_count: int  # 占位符数量


@dataclass
class TemplateFormatting:
    """模板格式信息"""
    font_requirements: Dict[str, Any]  # 字体要求
    paragraph_styles: Dict[str, Any]  # 段落样式
    spacing_rules: Dict[str, Any]  # 间距规范
    alignment_rules: Dict[str, Any]  # 对齐规则


@dataclass
class ContentRequirements:
    """内容要求信息"""
    word_limits: Dict[str, Any]  # 字数限制
    style_guide: Dict[str, Any]  # 风格指南
    terminology: Dict[str, Any]  # 术语规范
    structure_requirements: List[str]  # 结构要求


@dataclass
class FigureRequirements:
    """图表要求信息"""
    formats: List[str]  # 支持的图表格式
    numbering_rules: Dict[str, Any]  # 编号规则
    caption_format: Dict[str, Any]  # 标题格式
    placement_rules: Dict[str, Any]  # 放置规则


@dataclass
class TemplateIntelligence:
    """智能分析结果"""
    complexity_score: float  # 复杂度评分 (0-1)
    applicable_domains: List[str]  # 适用技术领域
    quality_score: float  # 质量评分 (0-1)
    suggestions: List[str]  # 改进建议
    template_type: str  # 模板类型
    completeness_score: float  # 完整性评分


@dataclass
class TemplateAnalysis:
    """完整的模板分析结果"""
    template_id: str
    template_name: str
    structure: TemplateStructure
    formatting: TemplateFormatting
    content_requirements: ContentRequirements
    figure_requirements: FigureRequirements
    intelligence: TemplateIntelligence
    analyzed_at: float
    file_path: str


class TemplateAnalyzer:
    """模板智能分析器"""

    def __init__(self):
        # 预定义的专利章节模式
        self.patent_section_patterns = {
            '标题': r'(标题|发明名称|专利名称|专利标题)',
            '技术领域': r'(技术领域|技术领域背景|所属技术领域)',
            '背景技术': r'(背景技术|现有技术|相关技术|技术背景)',
            '发明内容': r'(发明内容|技术方案|技术概述|发明概述)',
            '附图说明': r'(附图说明|图示说明|图表说明|图面说明)',
            '具体实施方式': r'(具体实施方式|实施例|具体实施例|实施方式)',
            '权利要求书': r'(权利要求书|权利要求|权项|专利权项)',
            '摘要': r'(摘要|技术摘要|内容摘要|发明摘要)'
        }

        # 占位符模式
        self.placeholder_patterns = [
            r'\{\{\s*([^}]+)\s*\}\}',  # {{placeholder}}
            r'\{\s*([^}]+)\s*\}',      # {placeholder}
            r'<\s*([^>]+)\s*>',        # <placeholder>
            r'\[\s*([^\]]+)\s*\]'      # [placeholder]
        ]

        # 字体大小映射
        self.font_size_mapping = {
            '初号': 42, '小初': 36, '一号': 26, '小一': 24,
            '二号': 22, '小二': 18, '三号': 16, '小三': 15,
            '四号': 14, '小四': 12, '五号': 10.5, '小五': 9
        }

        # 技术领域关键词
        self.domain_keywords = {
            '计算机软件': ['软件', '程序', '算法', '数据', '系统', '平台', '应用'],
            '电子通信': ['通信', '网络', '信号', '电子', '电路', '无线', '天线'],
            '机械制造': ['机械', '装置', '设备', '机构', '传动', '制造', '加工'],
            '化学材料': ['化学', '材料', '化合物', '合成', '反应', '组合物'],
            '医疗器械': ['医疗', '医用', '治疗', '诊断', '药物', '手术', '康复'],
            '新能源': ['能源', '电池', '光伏', '风能', '储能', '发电', '节能'],
            '人工智能': ['人工智能', 'AI', '机器学习', '深度学习', '神经网络', '算法']
        }

    def analyze_template(self, template_path: str, template_id: str, template_name: str) -> TemplateAnalysis:
        """
        分析模板文件并返回完整的分析结果

        Args:
            template_path: 模板文件路径
            template_id: 模板ID
            template_name: 模板名称

        Returns:
            TemplateAnalysis: 完整的模板分析结果
        """
        import time
        start_time = time.time()

        try:
            logger.info(f"开始分析模板: {template_name} ({template_id})")

            # 加载文档
            doc = Document(template_path)

            # 分析各个维度
            structure = self._analyze_structure(doc)
            formatting = self._analyze_formatting(doc)
            content_requirements = self._analyze_content_requirements(doc, template_name)
            figure_requirements = self._analyze_figure_requirements(doc)
            intelligence = self._analyze_intelligence(doc, structure, content_requirements)

            # 创建分析结果
            analysis = TemplateAnalysis(
                template_id=template_id,
                template_name=template_name,
                structure=structure,
                formatting=formatting,
                content_requirements=content_requirements,
                figure_requirements=figure_requirements,
                intelligence=intelligence,
                analyzed_at=time.time(),
                file_path=template_path
            )

            analysis_time = time.time() - start_time
            logger.info(f"模板分析完成: {template_name}, 耗时: {analysis_time:.2f}秒")

            return analysis

        except Exception as e:
            logger.error(f"模板分析失败: {template_name}, 错误: {str(e)}")
            raise

    def _analyze_structure(self, doc: Document) -> TemplateStructure:
        """分析文档结构"""
        hierarchy = {}
        sections = []
        placeholders = []
        max_level = 0

        current_section = None
        section_counter = 0

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            # 检查是否为标题
            heading_level = self._get_heading_level(paragraph)

            if heading_level > 0:
                # 这是一个标题
                section_counter += 1
                section_id = f"section_{section_counter}"

                # 识别章节类型
                section_type = self._identify_section_type(text)

                # 创建章节信息
                section_info = {
                    'id': section_id,
                    'title': text,
                    'level': heading_level,
                    'type': section_type,
                    'paragraph_count': 0,
                    'word_count': len(text),
                    'has_placeholder': bool(self._extract_placeholders(text))
                }

                sections.append(section_info)
                current_section = section_info

                # 更新层级结构
                if section_type not in hierarchy:
                    hierarchy[section_type] = []
                hierarchy[section_type].append(section_info)

                max_level = max(max_level, heading_level)

            elif current_section:
                # 这是正文内容
                current_section['paragraph_count'] += 1
                current_section['word_count'] += len(text)

                # 提取占位符
                paragraph_placeholders = self._extract_placeholders(text)
                placeholders.extend(paragraph_placeholders)

        # 统计所有占位符（包括标题中的）
        all_placeholders = []
        for section in sections:
            section_placeholders = self._extract_placeholders(section['title'])
            all_placeholders.extend(section_placeholders)
        all_placeholders.extend(placeholders)

        return TemplateStructure(
            hierarchy=hierarchy,
            sections=sections,
            section_count=len(sections),
            max_heading_level=max_level,
            placeholder_count=len(set(all_placeholders))
        )

    def _analyze_formatting(self, doc: Document) -> TemplateFormatting:
        """分析格式要求"""
        font_requirements = {}
        paragraph_styles = {}
        spacing_rules = {}
        alignment_rules = {}

        font_usage = {}
        style_usage = {}
        alignment_usage = {}

        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue

            # 分析字体
            if paragraph.runs:
                run = paragraph.runs[0]
                font_name = run.font.name or '默认'
                font_size = run.font.size

                if font_size:
                    # 转换为磅值
                    font_size_pt = font_size.pt if hasattr(font_size, 'pt') else font_size
                else:
                    font_size_pt = 12  # 默认大小

                font_key = f"{font_name}_{font_size_pt}pt"
                font_usage[font_key] = font_usage.get(font_key, 0) + 1

                # 分析样式
                style_name = paragraph.style.name if paragraph.style else 'Normal'
                style_usage[style_name] = style_usage.get(style_name, 0) + 1

            # 分析对齐
            alignment = paragraph.alignment
            if alignment:
                alignment_name = {
                    WD_PARAGRAPH_ALIGNMENT.LEFT: '左对齐',
                    WD_PARAGRAPH_ALIGNMENT.CENTER: '居中',
                    WD_PARAGRAPH_ALIGNMENT.RIGHT: '右对齐',
                    WD_PARAGRAPH_ALIGNMENT.JUSTIFY: '两端对齐',
                    WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE: '分散对齐'
                }.get(alignment, '未知')
                alignment_usage[alignment_name] = alignment_usage.get(alignment_name, 0) + 1

        # 提取主要要求
        if font_usage:
            most_common_font = max(font_usage, key=font_usage.get)
            font_requirements = {
                'primary_font': most_common_font,
                'font_variations': list(font_usage.keys()),
                'consistency_score': len(font_usage) / max(font_usage.values())
            }

        if style_usage:
            most_common_style = max(style_usage, key=style_usage.get)
            paragraph_styles = {
                'primary_style': most_common_style,
                'style_variations': list(style_usage.keys()),
                'style_distribution': style_usage
            }

        if alignment_usage:
            most_common_alignment = max(alignment_usage, key=alignment_usage.get)
            alignment_rules = {
                'primary_alignment': most_common_alignment,
                'alignment_distribution': alignment_usage
            }

        return TemplateFormatting(
            font_requirements=font_requirements,
            paragraph_styles=paragraph_styles,
            spacing_rules=spacing_rules,
            alignment_rules=alignment_rules
        )

    def _analyze_content_requirements(self, doc: Document, template_name: str) -> ContentRequirements:
        """分析内容要求"""
        word_limits = {}
        style_guide = {}
        terminology = {}
        structure_requirements = []

        # 从模板名称和内容推断要求
        all_text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())

        # 分析字数要求（从模板内容中提取）
        word_limit_patterns = [
            r'(\d+)[-~]?(\d+)?[个字]',
            r'(\d+)[-~]?(\d+)?字',
            r'不少于(\d+)字',
            r'不超过(\d+)字'
        ]

        for pattern in word_limit_patterns:
            matches = re.finditer(pattern, all_text)
            for match in matches:
                if match.group(2):
                    # 范围
                    word_limits['range'] = f"{match.group(1)}-{match.group(2)}字"
                else:
                    # 单一值
                    word_limits['limit'] = f"{match.group(1)}字"

        # 分析风格指南
        style_indicators = {
            '正式': ['正式', '规范', '标准', '严格'],
            '简洁': ['简洁', '精炼', '简明', '概括'],
            '详细': ['详细', '具体', '详尽', '完整'],
            '技术性': ['技术', '专业', '术语', '规范']
        }

        for style, keywords in style_indicators.items():
            count = sum(all_text.count(keyword) for keyword in keywords)
            if count > 0:
                style_guide[style] = count

        # 分析术语要求
        technical_terms = re.findall(r'[专业|术语|标准|规范|规定]', all_text)
        if technical_terms:
            terminology['has_terminology_requirements'] = True
            terminology['term_frequency'] = len(technical_terms)

        # 分析结构要求
        if '图' in all_text or '附图' in all_text:
            structure_requirements.append('包含图表或附图')
        if '实施例' in all_text or '具体实施' in all_text:
            structure_requirements.append('包含具体实施例')
        if '权利要求' in all_text:
            structure_requirements.append('包含权利要求书')

        return ContentRequirements(
            word_limits=word_limits,
            style_guide=style_guide,
            terminology=terminology,
            structure_requirements=structure_requirements
        )

    def _analyze_figure_requirements(self, doc: Document) -> FigureRequirements:
        """分析图表要求"""
        all_text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())

        formats = []
        numbering_rules = {}
        caption_format = {}
        placement_rules = {}

        # 分析支持的图表格式
        format_keywords = {
            '流程图': ['流程图', '流程', 'flow'],
            '结构图': ['结构图', '结构', 'structure'],
            '示意图': ['示意图', '示意', 'diagram'],
            '框图': ['框图', '框图', 'block'],
            '表格': ['表格', '表', 'table']
        }

        for format_type, keywords in format_keywords.items():
            if any(keyword in all_text for keyword in keywords):
                formats.append(format_type)

        # 分析编号规则
        numbering_patterns = [
            (r'图\s*(\d+)', '数字编号'),
            (r'图\s*(\d+)[-.](\d+)', '层级编号'),
            (r'Figure\s*(\d+)', '英文编号'),
            (r'Fig\.?\s*(\d+)', '缩写编号')
        ]

        for pattern, rule_type in numbering_patterns:
            if re.search(pattern, all_text):
                numbering_rules[rule_type] = True

        # 分析标题格式
        if '图' in all_text and ('标题' in all_text or '说明' in all_text):
            caption_format = {
                'requires_caption': True,
                'format': '下方说明'
            }

        # 分析放置规则
        if '正文' in all_text and '图' in all_text:
            placement_rules = {
                'placement': '文中插图',
                'reference_required': True
            }

        return FigureRequirements(
            formats=formats,
            numbering_rules=numbering_rules,
            caption_format=caption_format,
            placement_rules=placement_rules
        )

    def _analyze_intelligence(self, doc: Document, structure: TemplateStructure,
                            content_req: ContentRequirements) -> TemplateIntelligence:
        """智能分析"""
        complexity_score = 0.0
        applicable_domains = []
        quality_score = 0.0
        suggestions = []

        # 计算复杂度评分
        factors = []

        # 结构复杂度
        structure_complexity = min(structure.section_count / 8.0, 1.0)  # 标准化到0-1
        factors.append(('structure', structure_complexity))

        # 占位符复杂度
        placeholder_complexity = min(structure.placeholder_count / 10.0, 1.0)
        factors.append(('placeholder', placeholder_complexity))

        # 层级复杂度
        hierarchy_complexity = min(structure.max_heading_level / 4.0, 1.0)
        factors.append(('hierarchy', hierarchy_complexity))

        # 格式复杂度（基于章节类型多样性）
        section_types = len(set(s['type'] for s in structure.sections))
        format_complexity = min(section_types / 5.0, 1.0)
        factors.append(('format', format_complexity))

        # 计算加权平均复杂度
        weights = {'structure': 0.3, 'placeholder': 0.2, 'hierarchy': 0.2, 'format': 0.3}
        complexity_score = sum(score * weights.get(name, 0.2) for name, score in factors)

        # 识别适用领域
        all_text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        for domain, keywords in self.domain_keywords.items():
            domain_score = sum(all_text.count(keyword) for keyword in keywords)
            if domain_score > 0:
                applicable_domains.append(domain)

        # 计算质量评分
        quality_factors = []

        # 完整性评分
        required_sections = ['标题', '技术领域', '背景技术', '发明内容', '权利要求书', '摘要']
        found_sections = [s['type'] for s in structure.sections]
        completeness = len(set(required_sections) & set(found_sections)) / len(required_sections)
        quality_factors.append(('completeness', completeness))

        # 标准化评分
        standard_patterns = ['权利要求', '附图说明', '具体实施方式']
        standard_score = sum(1 for pattern in standard_patterns if pattern in all_text) / len(standard_patterns)
        quality_factors.append(('standardization', standard_score))

        # 清晰度评分（基于占位符使用）
        clarity_score = min(structure.placeholder_count / 5.0, 1.0) if structure.placeholder_count > 0 else 0.5
        quality_factors.append(('clarity', clarity_score))

        # 计算质量评分
        quality_weights = {'completeness': 0.4, 'standardization': 0.4, 'clarity': 0.2}
        quality_score = sum(score * weights.get(name, 0.33) for name, score in quality_factors)

        # 生成改进建议
        if completeness < 0.8:
            missing_sections = set(required_sections) - set(found_sections)
            suggestions.append(f"建议添加缺失的标准章节: {', '.join(missing_sections)}")

        if structure.placeholder_count == 0:
            suggestions.append("建议添加占位符以便更好地指导内容生成")

        if len(applicable_domains) == 0:
            suggestions.append("建议在模板中明确适用的技术领域")

        if structure.max_heading_level > 4:
            suggestions.append("标题层级过深，建议简化结构以提高可读性")

        # 确定模板类型
        if '发明' in all_text or '实用新型' in all_text:
            template_type = '发明专利模板'
        elif '外观' in all_text:
            template_type = '外观设计模板'
        else:
            template_type = '通用专利模板'

        return TemplateIntelligence(
            complexity_score=complexity_score,
            applicable_domains=applicable_domains,
            quality_score=quality_score,
            suggestions=suggestions,
            template_type=template_type,
            completeness_score=completeness
        )

    def _get_heading_level(self, paragraph) -> int:
        """获取段落的标题级别"""
        if paragraph.style.name.startswith('Heading'):
            try:
                return int(paragraph.style.name.split()[-1])
            except (ValueError, IndexError):
                pass

        # 基于字体大小判断
        if paragraph.runs:
            max_font_size = 0
            for run in paragraph.runs:
                if run.font.size:
                    size = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size
                    max_font_size = max(max_font_size, size)

            if max_font_size >= 20:
                return 1
            elif max_font_size >= 16:
                return 2
            elif max_font_size >= 14:
                return 3
            elif max_font_size >= 12:
                return 4

        return 0  # 不是标题

    def _identify_section_type(self, title: str) -> str:
        """识别章节类型"""
        title_lower = title.lower()

        for section_type, pattern in self.patent_section_patterns.items():
            if re.search(pattern, title, re.IGNORECASE):
                return section_type

        return '其他章节'

    def _extract_placeholders(self, text: str) -> List[str]:
        """提取文本中的占位符"""
        placeholders = []

        for pattern in self.placeholder_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                placeholder = match.group(1).strip()
                if placeholder and placeholder not in placeholders:
                    placeholders.append(placeholder)

        return placeholders