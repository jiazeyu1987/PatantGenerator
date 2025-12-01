"""
DOCX 文档生成器

负责将 Markdown 格式的专利文档转换为符合模板格式的 DOCX 文档。
"""

import re
import logging
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)


class MarkdownParser:
    """Markdown 解析器"""

    def __init__(self, content: str):
        self.content = content
        self.sections = {}
        self.metadata = {}
        self._parse_content()

    def _parse_content(self):
        """解析 Markdown 内容"""
        lines = self.content.split('\n')
        current_section = None
        current_content = []

        for line in lines:
            line = line.rstrip()

            # 检查是否为标题
            if line.startswith('#'):
                # 保存当前章节内容
                if current_section and current_content:
                    self.sections[current_section] = '\n'.join(current_content).strip()

                # 解析新章节
                current_section = self._parse_heading(line)
                current_content = []
            elif line.strip() == '':
                # 空行
                current_content.append('')
            else:
                current_content.append(line)

        # 保存最后一个章节
        if current_section and current_content:
            self.sections[current_section] = '\n'.join(current_content).strip()

    def _parse_heading(self, line: str) -> str:
        """解析标题行"""
        # 移除 # 号并清理
        title = re.sub(r'^#+\s*', '', line).strip()
        return title

    def get_section_content(self, section_name: str) -> Optional[str]:
        """获取指定章节的内容"""
        # 尝试精确匹配
        if section_name in self.sections:
            return self.sections[section_name]

        # 尝试模糊匹配
        for section, content in self.sections.items():
            if section_name.lower() in section.lower() or section.lower() in section_name.lower():
                return content

        return None

    def get_all_sections(self) -> Dict[str, str]:
        """获取所有章节"""
        return self.sections.copy()


class DOCXGenerator:
    """DOCX 文档生成器"""

    def __init__(self, template_path: str):
        """
        初始化生成器

        Args:
            template_path: 模板文件路径
        """
        self.template_path = template_path
        self.parser: Optional[MarkdownParser] = None

    def generate_from_markdown(self, markdown_content: str, output_path: str) -> bool:
        """
        从 Markdown 生成 DOCX 文档

        Args:
            markdown_content: Markdown 格式的专利内容
            output_path: 输出文件路径

        Returns:
            是否生成成功
        """
        try:
            # 解析 Markdown 内容
            self.parser = MarkdownParser(markdown_content)

            # 加载模板
            doc = Document(self.template_path)

            # 内容注入
            self._inject_content(doc)

            # 保存文档
            doc.save(output_path)

            logger.info(f"DOCX 文档生成成功: {output_path}")
            return True

        except Exception as e:
            logger.error(f"生成 DOCX 文档失败: {e}")
            return False

    def _inject_content(self, doc: Document):
        """将内容注入到模板中"""
        if not self.parser:
            raise ValueError("Markdown 解析器未初始化")

        # 获取所有章节
        sections = self.parser.get_all_sections()

        # 定义章节映射关系
        section_mapping = {
            # 标题类
            '标题': ['标题', '发明名称', '专利名称', '专利标题'],
            '技术领域': ['技术领域', '技术领域背景'],
            '背景技术': ['背景技术', '现有技术', '相关技术'],
            '发明内容': ['发明内容', '技术方案', '技术概述'],
            '附图说明': ['附图说明', '图示说明', '图表说明'],
            '具体实施方式': ['具体实施方式', '实施例', '具体实施例'],
            '权利要求书': ['权利要求书', '权利要求', '权项'],
            '摘要': ['摘要', '技术摘要', '内容摘要']
        }

        # 尝试找到并替换模板中的占位符
        self._replace_placeholders(doc, sections)

        # 尝试按章节名称匹配
        self._match_sections(doc, sections, section_mapping)

    def _replace_placeholders(self, doc: Document, sections: Dict[str, str]):
        """替换模板中的占位符"""
        placeholder_patterns = [
            r'\{\{\s*([^}]+)\s*\}\}',  # {{placeholder}}
            r'\{\s*([^}]+)\s*\}',      # {placeholder}
            r'<\s*([^>]+)\s*>',        # <placeholder>
            r'\[\s*([^\]]+)\s*\]',      # [placeholder]
        ]

        # 在段落中查找和替换占位符
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            new_text = original_text

            for pattern in placeholder_patterns:
                matches = re.finditer(pattern, new_text)
                for match in matches:
                    placeholder = match.group(1).strip()

                    # 尝试在章节中找到对应内容
                    content = self._find_content_for_placeholder(placeholder, sections)
                    if content:
                        new_text = new_text.replace(match.group(0), content)

            if new_text != original_text:
                paragraph.text = new_text

    def _find_content_for_placeholder(self, placeholder: str, sections: Dict[str, str]) -> Optional[str]:
        """为占位符查找对应的内容"""
        # 直接匹配章节名称
        if placeholder in sections:
            return sections[placeholder]

        # 模糊匹配
        placeholder_lower = placeholder.lower()

        for section_name, content in sections.items():
            section_lower = section_name.lower()

            # 检查占位符是否是章节名称的一部分
            if placeholder_lower in section_lower or section_lower in placeholder_lower:
                return content

            # 检查关键词匹配
            keywords = ['标题', '领域', '背景', '内容', '说明', '方式', '要求', '摘要']
            for keyword in keywords:
                if keyword in placeholder_lower and keyword in section_lower:
                    return content

        return None

    def _match_sections(self, doc: Document, sections: Dict[str, str], section_mapping: Dict[str, List[str]]):
        """按章节名称匹配内容"""
        for target_section, possible_names in section_mapping.items():
            content = None

            # 尝试多种可能的名称
            for name in possible_names:
                if name in sections:
                    content = sections[name]
                    break
                else:
                    # 模糊匹配
                    for section_name in sections.keys():
                        if name.lower() in section_name.lower() or section_name.lower() in name.lower():
                            content = sections[section_name]
                            break
                if content:
                    break

            # 如果找到内容，尝试注入到文档中
            if content:
                self._inject_section_content(doc, target_section, content)

    def _inject_section_content(self, doc: Document, section_name: str, content: str):
        """将章节内容注入到文档中"""
        try:
            # 查找包含章节名称的段落
            for paragraph in doc.paragraphs:
                if section_name in paragraph.text:
                    # 找到目标段落，在其后插入内容
                    self._insert_content_after_paragraph(doc, paragraph, content)
                    break
        except Exception as e:
            logger.warning(f"注入章节内容失败 {section_name}: {e}")

    def _insert_content_after_paragraph(self, doc: Document, target_paragraph, content: str):
        """在指定段落后插入内容"""
        try:
            # 获取目标段落的父元素
            parent = target_paragraph._element.getparent()

            # 获取目标段落的索引
            target_index = parent.index(target_paragraph._element)

            # 分割内容为行
            content_lines = content.split('\n')

            # 在目标位置插入新段落
            for i, line in enumerate(content_lines):
                if line.strip():  # 跳过空行
                    new_paragraph = target_paragraph._p.add_p().add_t(line)
                    new_paragraph = parent.insert(target_index + i + 1, new_paragraph)

        except Exception as e:
            logger.warning(f"插入内容失败: {e}")

    @staticmethod
    def validate_template(template_path: str) -> Tuple[bool, str]:
        """验证模板文件"""
        try:
            doc = Document(template_path)

            # 基本检查：是否有内容
            if len(doc.paragraphs) == 0:
                return False, "模板文件为空"

            # 检查是否包含专利相关章节的标识
            content = '\n'.join(p.text for p in doc.paragraphs)
            patent_keywords = [
                '发明', '专利', '权利要求', '摘要',
                '技术方案', '背景技术', '具体实施'
            ]

            has_patent_content = any(keyword in content for keyword in patent_keywords)

            if not has_patent_content:
                return False, "模板不包含专利相关内容标识"

            return True, "模板验证通过"

        except Exception as e:
            return False, f"模板验证失败: {str(e)}"


def generate_patent_docx(
    markdown_content: str,
    template_path: str,
    output_path: str
) -> bool:
    """
    生成专利 DOCX 文档

    Args:
        markdown_content: Markdown 格式的专利内容
        template_path: 模板文件路径
        output_path: 输出文件路径

    Returns:
        是否生成成功
    """
    try:
        generator = DOCXGenerator(template_path)
        return generator.generate_from_markdown(markdown_content, output_path)
    except Exception as e:
        logger.error(f"生成专利 DOCX 文档失败: {e}")
        return False


def validate_patent_template(template_path: str) -> Tuple[bool, str]:
    """
    验证专利模板

    Args:
        template_path: 模板文件路径

    Returns:
        (是否有效, 验证消息)
    """
    return DOCXGenerator.validate_template(template_path)