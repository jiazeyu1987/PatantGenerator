"""
创建示例专利模板文件
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os

def create_basic_template():
    """创建基础的专利模板"""
    doc = Document()

    # 设置字体
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(12)

    # 添加标题
    title = doc.add_heading('专利标题', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # 空行

    # 技术领域
    tech_field = doc.add_heading('技术领域', level=2)
    tech_field.add_run('\n').add_paragraph('本发明涉及技术领域，具体是一种...')

    doc.add_paragraph()  # 空行

    # 背景技术
    background = doc.add_heading('背景技术', level=2)
    background.add_run('\n').add_paragraph('现有技术中存在...等问题。')

    doc.add_paragraph()  # 空行

    # 发明内容
    content = doc.add_heading('发明内容', level=2)
    content.add_run('\n').add_paragraph('本发明的目的是提供一种...，以解决现有技术中存在的问题。')

    doc.add_paragraph()  # 空行

    # 附图说明
    figure_desc = doc.add_heading('附图说明', level=2)
    figure_desc.add_run('\n').add_paragraph('图1是本发明实施例的结构示意图。')

    doc.add_paragraph()  # 空行

    # 具体实施方式
    implementation = doc.add_heading('具体实施方式', level=2)
    implementation.add_run('\n').add_paragraph('下面结合附图对本发明的具体实施方式进行详细描述。')

    doc.add_paragraph()  # 空行

    # 权利要求书
    claims = doc.add_heading('权利要求书', level=2)
    claims.add_run('\n').add_paragraph('1. 一种...，其特征在于包括：...')

    doc.add_paragraph()  # 空行

    # 摘要
    abstract = doc.add_heading('摘要', level=2)
    abstract.add_run('\n').add_paragraph('本发明公开了一种...，具有...等优点。')

    # 保存模板
    template_path = 'backend/templates_store/基础专利模板.docx'
    os.makedirs(os.path.dirname(template_path), exist_ok=True)
    doc.save(template_path)

    print(f"基础专利模板已创建: {template_path}")

def create_advanced_template():
    """创建高级专利模板，包含占位符"""
    doc = Document()

    # 设置字体
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(12)

    # 使用占位符的模板
    title = doc.add_heading('{{标题}}', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()

    # 技术领域
    doc.add_heading('技术领域', level=2)
    doc.add_paragraph('{{技术领域}}')

    doc.add_paragraph()

    # 背景技术
    doc.add_heading('背景技术', level=2)
    doc.add_paragraph('{{背景技术}}')

    doc.add_paragraph()

    # 发明内容
    doc.add_heading('发明内容', level=2)
    doc.add_paragraph('{{发明内容}}')

    doc.add_paragraph()

    # 附图说明
    doc.add_heading('附图说明', level=2)
    doc.add_paragraph('{{附图说明}}')

    doc.add_paragraph()

    # 具体实施方式
    doc.add_heading('具体实施方式', level=2)
    doc.add_paragraph('{{具体实施方式}}')

    doc.add_paragraph()

    # 权利要求书
    doc.add_heading('权利要求书', level=2)
    doc.add_paragraph('{{权利要求书}}')

    doc.add_paragraph()

    # 摘要
    doc.add_heading('摘要', level=2)
    doc.add_paragraph('{{摘要}}')

    # 保存模板
    template_path = 'backend/templates_store/高级专利模板（占位符）.docx'
    os.makedirs(os.path.dirname(template_path), exist_ok=True)
    doc.save(template_path)

    print(f"高级专利模板已创建: {template_path}")

if __name__ == "__main__":
    create_basic_template()
    create_advanced_template()
    print("\n示例模板创建完成！")
    print("请在后端服务器中测试模板功能。")