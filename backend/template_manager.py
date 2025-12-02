"""
模板管理器

负责管理和加载专利模板文件，支持模板扫描、验证、选择等功能。
"""

import os
import json
import logging
import time
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path
from datetime import datetime
import re

from docx_generator import validate_patent_template
from template_analyzer import TemplateAnalyzer, TemplateAnalysis

logger = logging.getLogger(__name__)


class TemplateInfo:
    """模板信息数据类"""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.file_name = Path(file_path).name
        self.template_id = self._generate_id()
        self.name = self._extract_name()
        self.description = ""
        self.created_at = None
        self.modified_at = None
        self.file_size = 0
        self.is_valid = False
        self.placeholder_count = 0
        self.sections = []

        # 深度分析结果
        self.analysis: Optional[TemplateAnalysis] = None
        self.analysis_timestamp: Optional[float] = None
        self.analysis_cached: bool = False

        self._load_file_info()
        self._analyze_template()

    def _generate_id(self) -> str:
        """生成模板ID"""
        # 使用文件名（不含扩展名）作为ID
        base_name = Path(self.file_name).stem
        # 转换为小写并用下划线替换空格和特殊字符
        return re.sub(r'[^\w]', '_', base_name.lower())

    def _extract_name(self) -> str:
        """从文件名提取模板名称"""
        # 移除扩展名，用空格替换下划线
        base_name = Path(self.file_name).stem
        name = re.sub(r'_+', ' ', base_name)
        return name.title()

    def _load_file_info(self):
        """加载文件基本信息"""
        try:
            path = Path(self.file_path)
            if path.exists():
                stat = path.stat()
                self.created_at = datetime.fromtimestamp(stat.st_ctime)
                self.modified_at = datetime.fromtimestamp(stat.st_mtime)
                self.file_size = stat.st_size
        except Exception as e:
            logger.error(f"加载模板文件信息失败 {self.file_path}: {e}")

    def _analyze_template(self):
        """分析模板内容"""
        try:
            from docx import Document

            doc = Document(self.file_path)
            self.placeholder_count = self._count_placeholders(doc)
            self.sections = self._extract_sections(doc)
            self.is_valid = self._validate_template(doc)

        except Exception as e:
            logger.error(f"分析模板失败 {self.file_path}: {e}")
            self.is_valid = False

    def _count_placeholders(self, doc) -> int:
        """统计模板中的占位符数量"""
        count = 0
        placeholder_patterns = [
            r'\{\{[^}]+\}\}',  # {{placeholder}}
            r'\{[^}]+\}',      # {placeholder}
            r'<[^>]+>',        # <placeholder>
            r'\[[^\]]+\]',      # [placeholder]
        ]

        for paragraph in doc.paragraphs:
            text = paragraph.text
            for pattern in placeholder_patterns:
                matches = re.findall(pattern, text)
                count += len(matches)

        return count

    def _extract_sections(self, doc) -> List[str]:
        """提取模板中的章节标题"""
        sections = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text and self._is_heading(text):
                sections.append(text)
        return sections

    def _is_heading(self, text: str) -> bool:
        """判断是否为标题"""
        # 简单的标题判断逻辑
        heading_indicators = [
            '第', '章', '节', '部分',
            '标题', '题目', '名称',
            '技术领域', '背景技术', '发明内容',
            '附图说明', '具体实施方式', '权利要求书', '摘要'
        ]
        return any(indicator in text for indicator in heading_indicators)

    def _validate_template(self, doc) -> bool:
        """验证模板有效性"""
        try:
            # 检查文档是否有内容
            if len(doc.paragraphs) == 0:
                return False

            # 检查是否包含基本的专利章节
            patent_sections = [
                '发明内容', '技术方案', '背景技术',
                '具体实施方式', '权利要求书', '摘要'
            ]

            content = '\n'.join(p.text for p in doc.paragraphs)
            has_patent_content = any(section in content for section in patent_sections)

            return has_patent_content or self.placeholder_count > 0

        except Exception:
            return False

    def to_dict(self, include_analysis: bool = False) -> Dict[str, Any]:
        """转换为字典格式"""
        result = {
            'id': self.template_id,
            'name': self.name,
            'description': self.description,
            'file_name': self.file_name,
            'file_path': self.file_path,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'modified_at': self.modified_at.isoformat() if self.modified_at else None,
            'file_size': self.file_size,
            'is_valid': self.is_valid,
            'placeholder_count': self.placeholder_count,
            'sections': self.sections,
            'has_analysis': self.analysis is not None,
            'analysis_timestamp': self.analysis_timestamp
        }

        # 包含深度分析结果
        if include_analysis and self.analysis:
            result['analysis'] = {
                'structure': {
                    'section_count': self.analysis.structure.section_count,
                    'max_heading_level': self.analysis.structure.max_heading_level,
                    'placeholder_count': self.analysis.structure.placeholder_count,
                    'sections': [{'id': s['id'], 'title': s['title'], 'level': s['level'], 'type': s['type']}
                               for s in self.analysis.structure.sections]
                },
                'formatting': {
                    'font_requirements': self.analysis.formatting.font_requirements,
                    'paragraph_styles': self.analysis.formatting.paragraph_styles,
                    'alignment_rules': self.analysis.formatting.alignment_rules
                },
                'content_requirements': {
                    'word_limits': self.analysis.content_requirements.word_limits,
                    'style_guide': self.analysis.content_requirements.style_guide,
                    'structure_requirements': self.analysis.content_requirements.structure_requirements
                },
                'figure_requirements': {
                    'formats': self.analysis.figure_requirements.formats,
                    'numbering_rules': self.analysis.figure_requirements.numbering_rules,
                    'caption_format': self.analysis.figure_requirements.caption_format
                },
                'intelligence': {
                    'complexity_score': round(self.analysis.intelligence.complexity_score, 3),
                    'quality_score': round(self.analysis.intelligence.quality_score, 3),
                    'completeness_score': round(self.analysis.intelligence.completeness_score, 3),
                    'template_type': self.analysis.intelligence.template_type,
                    'applicable_domains': self.analysis.intelligence.applicable_domains,
                    'suggestions': self.analysis.intelligence.suggestions
                }
            }

        return result


class TemplateManager:
    """模板管理器"""

    def __init__(self, template_dir: str = None):
        """
        初始化模板管理器

        Args:
            template_dir: 模板目录路径，默认为 backend/templates_store
        """
        if template_dir is None:
            # 默认模板目录
            current_dir = Path(__file__).parent
            template_dir = current_dir / "templates_store"

        self.template_dir = Path(template_dir)
        self.templates: Dict[str, TemplateInfo] = {}
        self.default_template_id: Optional[str] = None
        self.metadata_file = self.template_dir / "templates_metadata.json"
        self.analysis_cache_file = self.template_dir / "templates_analysis.json"

        # 初始化分析器
        self.analyzer = TemplateAnalyzer()

        # 分析缓存
        self._analysis_cache: Dict[str, Dict[str, Any]] = {}

        logger.info(f"模板管理器初始化，目录: {self.template_dir}")

        # 确保模板目录存在
        self._ensure_template_dir()

        # 加载所有模板
        self._load_templates()

        # 加载分析缓存
        self._load_analysis_cache()

    def _ensure_template_dir(self):
        """确保模板目录存在"""
        try:
            self.template_dir.mkdir(parents=True, exist_ok=True)
            logger.info(f"模板目录已确保存在: {self.template_dir}")
        except Exception as e:
            logger.error(f"创建模板目录失败: {e}")

    def _load_templates(self):
        """加载所有模板文件"""
        logger.info(f"🔍 开始加载模板，扫描目录: {self.template_dir}")

        if not self.template_dir.exists():
            logger.error(f"❌ 模板目录不存在: {self.template_dir}")
            return

        # 扫描所有 .docx 文件
        docx_files = list(self.template_dir.glob("*.docx"))
        logger.info(f"📁 找到 {len(docx_files)} 个 .docx 文件: {[f.name for f in docx_files]}")

        if not docx_files:
            logger.warning(f"⚠️ 未找到模板文件: {self.template_dir}")
            return

        loaded_count = 0
        for file_path in docx_files:
            try:
                logger.info(f"📋 正在加载模板: {file_path.name}")
                template_info = TemplateInfo(str(file_path))
                self.templates[template_info.template_id] = template_info
                loaded_count += 1
                logger.info(f"✅ 模板加载成功: {template_info.name} (ID: {template_info.template_id}, 大小: {template_info.file_size} bytes)")
            except Exception as e:
                logger.error(f"❌ 加载模板失败 {file_path}: {e}")
                import traceback
                logger.error(f"详细错误: {traceback.format_exc()}")

        # 设置默认模板（第一个有效的模板）
        self._set_default_template()

        logger.info(f"🎉 模板加载完成，共加载 {loaded_count} 个模板")
        logger.info(f"📋 可用模板ID: {list(self.templates.keys())}")
        if self.default_template_id:
            logger.info(f"🎯 默认模板ID: {self.default_template_id}")

    def _set_default_template(self):
        """设置默认模板"""
        for template_info in self.templates.values():
            if template_info.is_valid:
                self.default_template_id = template_info.template_id
                break

    def get_template_list(self) -> List[Dict[str, Any]]:
        """获取模板列表"""
        return [template.to_dict() for template in self.templates.values()]

    def get_template_info(self, template_id: str) -> Optional[Dict[str, Any]]:
        """获取指定模板信息"""
        template = self.templates.get(template_id)
        return template.to_dict() if template else None

    def get_default_template(self) -> Optional[Dict[str, Any]]:
        """获取默认模板"""
        if self.default_template_id:
            return self.get_template_info(self.default_template_id)
        return None

    def validate_template(self, template_path: str) -> Tuple[bool, str]:
        """验证模板文件"""
        try:
            if not Path(template_path).exists():
                return False, "模板文件不存在"

            template_info = TemplateInfo(template_path)

            if not template_info.is_valid:
                return False, "模板格式无效或不完整"

            return True, "模板验证通过"

        except Exception as e:
            return False, f"模板验证失败: {str(e)}"

    def get_template_content(self, template_id: str) -> Optional[Any]:
        """获取模板文档对象"""
        template = self.templates.get(template_id)
        if not template or not template.is_valid:
            return None

        try:
            from docx import Document
            return Document(template.file_path)
        except Exception as e:
            logger.error(f"加载模板文档失败 {template_id}: {e}")
            return None

    def save_template_metadata(self):
        """保存模板元数据"""
        try:
            metadata = {
                'templates': {
                    template_id: template.to_dict()
                    for template_id, template in self.templates.items()
                },
                'default_template_id': self.default_template_id,
                'last_updated': datetime.now().isoformat()
            }

            with open(self.metadata_file, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, ensure_ascii=False, indent=2)

            logger.debug("模板元数据已保存")

        except Exception as e:
            logger.error(f"保存模板元数据失败: {e}")

    def load_template_metadata(self):
        """加载模板元数据"""
        try:
            if not self.metadata_file.exists():
                return

            with open(self.metadata_file, 'r', encoding='utf-8') as f:
                metadata = json.load(f)

            # 加载模板描述信息
            templates_metadata = metadata.get('templates', {})
            for template_id, template_data in templates_metadata.items():
                if template_id in self.templates:
                    self.templates[template_id].description = template_data.get('description', '')

            logger.debug("模板元数据已加载")

        except Exception as e:
            logger.error(f"加载模板元数据失败: {e}")

    def reload_templates(self):
        """重新加载所有模板"""
        logger.info("重新加载模板")
        self.templates.clear()
        self.default_template_id = None
        self._load_templates()
        self.load_template_metadata()

    def get_template_analysis(self, template_id: str, force_reanalyze: bool = False) -> Optional[TemplateAnalysis]:
        """
        获取模板深度分析结果

        Args:
            template_id: 模板ID
            force_reanalyze: 是否强制重新分析

        Returns:
            TemplateAnalysis: 分析结果，如果失败返回None
        """
        template = self.templates.get(template_id)
        if not template or not template.is_valid:
            logger.warning(f"模板不存在或无效: {template_id}")
            return None

        # 检查缓存
        if not force_reanalyze and template.analysis and template.analysis_cached:
            # 检查文件是否被修改
            if template.modified_at and template.modified_at.timestamp() <= template.analysis_timestamp:
                logger.debug(f"使用缓存的分析结果: {template_id}")
                return template.analysis

        try:
            logger.info(f"开始深度分析模板: {template.name} ({template_id})")
            start_time = time.time()

            # 执行分析
            analysis = self.analyzer.analyze_template(
                template.file_path,
                template.template_id,
                template.name
            )

            # 缓存结果
            template.analysis = analysis
            template.analysis_timestamp = time.time()
            template.analysis_cached = True

            # 更新内存缓存
            self._analysis_cache[template_id] = {
                'analysis_data': analysis,
                'timestamp': template.analysis_timestamp,
                'file_mtime': template.modified_at.timestamp() if template.modified_at else 0
            }

            # 保存缓存到文件
            self._save_analysis_cache()

            analysis_time = time.time() - start_time
            logger.info(f"模板分析完成: {template.name}, 耗时: {analysis_time:.2f}秒")

            return analysis

        except Exception as e:
            logger.error(f"模板分析失败: {template.name}, 错误: {str(e)}")
            template.analysis_cached = False
            return None

    def get_template_analysis_summary(self, template_id: str) -> Optional[Dict[str, Any]]:
        """
        获取模板分析摘要（用于API返回）

        Args:
            template_id: 模板ID

        Returns:
            Dict: 分析摘要，如果失败返回None
        """
        analysis = self.get_template_analysis(template_id)
        if not analysis:
            return None

        return {
            'complexity_score': round(analysis.intelligence.complexity_score, 3),
            'quality_score': round(analysis.intelligence.quality_score, 3),
            'completeness_score': round(analysis.intelligence.completeness_score, 3),
            'template_type': analysis.intelligence.template_type,
            'applicable_domains': analysis.intelligence.applicable_domains,
            'section_count': analysis.structure.section_count,
            'placeholder_count': analysis.structure.placeholder_count,
            'suggestions_count': len(analysis.intelligence.suggestions),
            'analysis_timestamp': analysis.analyzed_at
        }

    def analyze_all_templates(self, force_reanalyze: bool = False) -> Dict[str, Any]:
        """
        分析所有模板

        Args:
            force_reanalyze: 是否强制重新分析所有模板

        Returns:
            Dict: 分析统计结果
        """
        logger.info("开始批量分析所有模板")
        start_time = time.time()

        results = {
            'total': len(self.templates),
            'analyzed': 0,
            'failed': 0,
            'skipped': 0,
            'details': {}
        }

        for template_id, template in self.templates.items():
            if not template.is_valid:
                results['skipped'] += 1
                continue

            try:
                analysis = self.get_template_analysis(template_id, force_reanalyze)
                if analysis:
                    results['analyzed'] += 1
                    results['details'][template_id] = {
                        'name': template.name,
                        'complexity_score': analysis.intelligence.complexity_score,
                        'quality_score': analysis.intelligence.quality_score,
                        'analysis_time': analysis.analyzed_at
                    }
                else:
                    results['failed'] += 1
            except Exception as e:
                logger.error(f"分析模板失败 {template_id}: {e}")
                results['failed'] += 1

        total_time = time.time() - start_time
        logger.info(f"批量分析完成: {results['analyzed']}/{results['total']} 成功, 耗时: {total_time:.2f}秒")

        results['total_time'] = total_time
        results['average_time'] = total_time / max(results['analyzed'], 1)

        return results

    def _load_analysis_cache(self):
        """加载分析缓存"""
        try:
            if not self.analysis_cache_file.exists():
                return

            with open(self.analysis_cache_file, 'r', encoding='utf-8') as f:
                cache_data = json.load(f)

            for template_id, cache_item in cache_data.items():
                if template_id in self.templates:
                    self._analysis_cache[template_id] = cache_item

                    # 检查缓存是否仍然有效
                    template = self.templates[template_id]
                    file_mtime = template.modified_at.timestamp() if template.modified_at else 0
                    cache_timestamp = cache_item.get('timestamp', 0)

                    if file_mtime <= cache_timestamp:
                        # 缓存有效，重建分析对象
                        try:
                            # 这里简化处理，实际使用时可以序列化/反序列化TemplateAnalysis对象
                            template.analysis_cached = True
                            template.analysis_timestamp = cache_timestamp
                        except Exception as e:
                            logger.warning(f"重建分析缓存失败 {template_id}: {e}")

            logger.debug(f"分析缓存已加载，共 {len(self._analysis_cache)} 个模板")

        except Exception as e:
            logger.error(f"加载分析缓存失败: {e}")

    def _save_analysis_cache(self):
        """保存分析缓存"""
        try:
            # 只保存基本信息，避免序列化复杂对象
            cache_data = {}
            for template_id, cache_item in self._analysis_cache.items():
                cache_data[template_id] = {
                    'timestamp': cache_item.get('timestamp', 0),
                    'file_mtime': cache_item.get('file_mtime', 0)
                }

            with open(self.analysis_cache_file, 'w', encoding='utf-8') as f:
                json.dump(cache_data, f, ensure_ascii=False, indent=2)

            logger.debug("分析缓存已保存")

        except Exception as e:
            logger.error(f"保存分析缓存失败: {e}")

    def get_analysis_stats(self) -> Dict[str, Any]:
        """获取分析统计信息"""
        analyzed_templates = sum(1 for t in self.templates.values() if t.analysis_cached)
        cached_templates = len(self._analysis_cache)

        return {
            'total_templates': len(self.templates),
            'analyzed_templates': analyzed_templates,
            'cached_templates': cached_templates,
            'analysis_coverage': round(analyzed_templates / max(len(self.templates), 1), 3),
            'cache_file_path': str(self.analysis_cache_file)
        }

    def get_stats(self) -> Dict[str, Any]:
        """获取模板统计信息"""
        total_templates = len(self.templates)
        valid_templates = sum(1 for t in self.templates.values() if t.is_valid)
        total_size = sum(t.file_size for t in self.templates.values())

        # 包含分析统计
        analysis_stats = self.get_analysis_stats()

        return {
            'total_templates': total_templates,
            'valid_templates': valid_templates,
            'invalid_templates': total_templates - valid_templates,
            'total_size_bytes': total_size,
            'total_size_mb': round(total_size / (1024 * 1024), 2),
            'default_template_id': self.default_template_id,
            'template_dir': str(self.template_dir),
            **analysis_stats
        }


# 全局模板管理器实例
_template_manager: Optional[TemplateManager] = None


def get_template_manager() -> TemplateManager:
    """获取全局模板管理器实例"""
    global _template_manager
    if _template_manager is None:
        _template_manager = TemplateManager()
    return _template_manager


def get_template_manager_with_dir(template_dir: str) -> TemplateManager:
    """获取指定目录的模板管理器实例"""
    return TemplateManager(template_dir)